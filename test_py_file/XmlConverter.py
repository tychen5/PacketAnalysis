#!/usr/bin/env python
# coding: utf-8

# In[ ]:


from docx import Document
import xml.etree.ElementTree as ET
from xml.dom import minidom
from pathlib import Path


# In[ ]:


class XmlConverter:
    def __init__(self, xmlFilePath, category):
        self.xml = self.readXmlFile(xmlFilePath)
    
    def readXmlFile(self, xmlFilePath):
        tree = ET.parse(xmlFilePath)
        root = tree.getroot()
        return root
    
    def prettify(self, elem):
        """Return a pretty-printed XML string for the Element.
        """
        rough_string = ET.tostring(elem, 'utf-8')
        reparsed = minidom.parseString(rough_string)
        return reparsed.toprettyxml(indent="  ")
    
    def convertLoner2Docx(self, outputDocxDir):
        Path(outputDocxDir).mkdir(parents=True, exist_ok=True)
        root = self.xml
        for loner in root:
            document = Document()
            document.add_heading('LonerIP Payload Analysis', 0)
            loner_id = loner[0].text
            number_of_clusters = loner[1].text
            number_of_protocols = loner[2].text
            src_ip= loner[3].text
            country = loner[4].text
            clusters = loner[5]

            document.add_heading('lonerIP ID: ' + loner_id, level=1)

            table = document.add_table(rows=1, cols=5, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'lonerIP ID'
            hdr_cells[1].text = '# clusters'
            hdr_cells[2].text = '# distinct protocols'
            hdr_cells[3].text = 'src.IP'
            hdr_cells[4].text = 'country'

            row_cells = table.add_row().cells
            row_cells[0].text = loner_id
            row_cells[1].text = number_of_clusters
            row_cells[2].text = number_of_protocols
            row_cells[3].text = src_ip
            row_cells[4].text = country
            
            for cluster in clusters:
                document.add_paragraph(cluster.text, style='List Bullet')
                for idx in range(len(cluster)):
                    document.add_paragraph(cluster[idx].text, style='List Bullet 2')
                    if idx == 9:
                        break
            outFilePath = outputDocxDir + "lonerIP_" + loner_id + ".docx"
            document.save(outFilePath)
#             break
    
    def convert2Docx(self, outputDocxDir):
        Path(outputDocxDir).mkdir(parents=True, exist_ok=True)
        root = self.xml
        for group in root:
            document = Document()
            document.add_heading('Group Payload Analysis', 0)
            group_id = group[0].text
            number_of_clusters = group[1].text
            number_of_protocols = group[2].text
            src_ips= group[3]
            number_of_ip = len(src_ips)
            countries = group[4]
            number_of_country = len(countries)
            clusters = group[5]

            document.add_heading('groupID: ' + loner_id, level=1)

            table = document.add_table(rows=1, cols=7, style='Table Grid')
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'groupID'
            hdr_cells[1].text = '# clusters'
            hdr_cells[2].text = '# distinct protocols'
            hdr_cells[3].text = '# of src.IP'
            hdr_cells[4].text = 'src.IP(s)'
            hdr_cells[5].text = '# countries'
            hdr_cells[6].text = 'country'

            row_cells = table.add_row().cells
            row_cells[0].text = group_id
            row_cells[1].text = number_of_clusters
            row_cells[2].text = number_of_protocols
            row_cells[3].text = str(number_of_ip)
            for i in range(len(src_ips)): # at most 10 IPs in docx
                ip = src_ips[i]
                row_cells[4].text += ip.text + ', '
                if i == 9:
                    row_cells[4].text += "..."
                    break
            row_cells[5].text = str(number_of_country)
            counter = 0
            for country in countries: # at most 10 countries in docx
                row_cells[6].text += country.text + ', '
                counter += 1
                if counter == 9:
                    row_cells[6].text += "..."
                    break
            
            for cluster in clusters: # at most 10 packet payload in docx
                document.add_paragraph(cluster.text, style='List Bullet')
                for idx in range(len(cluster)):
                    document.add_paragraph(cluster[idx].text, style='List Bullet 2')
                    if idx == 9:
                        break
            outFilePath = outputDocxDir + "group_" + group_id + ".docx"
            document.save(outFilePath)


# In[ ]:


converter = XmlConverter("./xml/top_correlate_0.1_loner_france.xml")
converter.convert2Docx("./payloadDocx/group_0.9_wo_tds/all/")


# In[ ]:


# converter.convertLoner2Docx("./payloadDocx/loner_0.1/france/")


# In[ ]:




